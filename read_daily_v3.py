"""
功能：将断面水质日报数据汇总至excel
需设定参数：
    dirpath：日报所在文件夹（无其他文件）
    outxlsx：输出文件路径
"""
import numpy as np
from docx import Document
import datetime
import pandas as pd
import os
import matplotlib.pyplot as plt
import matplotlib as mpl


def read_daily_report(path: str):
    document = Document(path)
    # 读取日期
    pl = [paragraph.text for paragraph in document.paragraphs]
    date = pl[1][pl[1].find('监测日期：') + 5:]
    d = datetime.datetime.strptime(date, '%Y年%m月%d日')
    date2 = d.strftime('%Y{}%m{}%d'.format('/', '/'))
    print(date2)

    def read_name(table, start_idx, end_idx):
        name = []
        for i in range(start_idx, len(table.rows) - end_idx):
            name.append(table.cell(i, 0).text)
        return name

    # 读取表格
    def read_table(table, start_idx, end_idx):
        col = []
        # print(len(table.columns))
        for i in range(start_idx, len(table.rows) - end_idx):
            col.append(table.cell(i, 2).text)
            col.append(table.cell(i, 3).text)
            col.append(table.cell(i, 4).text)
            # 日报的表格读取有时会出现重复列的情况（NH4），但直接打开显示无问题
            if len(table.columns) == 7:
                col.append(table.cell(i, 6).text)
            else:
                col.append(table.cell(i, 5).text)

        return col

    tables = document.tables
    print(f'''文件内共{len(tables)}个表格''')
    # 只需要第一个表格 国省考断面
    name0 = read_name(tables[0], 2, 2)
    # name1 = read_name(tables[1], 2, 2)
    # name2 = read_name(tables[2], 2, 2)
    # name3 = read_name(tables[3], 2, 3)
    col0 = np.array(read_table(tables[0], 2, 2)).reshape(-1, 4)
    # col1 = np.array(read_table(tables[1], 2, 2)).reshape(-1,4)
    # col2 = np.array(read_table(tables[2], 2, 2)).reshape(-1,4)
    # col3 = np.array(read_table(tables[3], 2, 3)).reshape(-1,4)
    # name = np.concatenate((name0,name1,name2,name3),axis=0)
    # result = np.concatenate((col0,col1,col2,col3),axis=0)

    name = name0
    result = col0
    return date2, name, result


def arrange_daily_report(dirpath: str, outxlsx: str, outxlsx1: str,show_days:int):
    # 遍历日报
    file_list = os.listdir(dirpath)
    file_list.sort(key=lambda x: (int(x.split('.')[1]), int(x.split('.')[2])))
    print(f"""日报个数为：{len(file_list)}""")

    # 读取日报
    O2 = {}
    CODM = {}
    NH4 = {}
    TP = {}
    for file in file_list:
        if os.path.splitext(file)[1] == '.docx':
            print(dirpath + '\\' + file)
            date, name, data = read_daily_report(dirpath + '\\' + file)
            O2[str(date)] = data[:, 0].tolist()
            CODM[str(date)] = data[:, 1].tolist()
            NH4[str(date)] = data[:, 2].tolist()
            TP[str(date)] = data[:, 3].tolist()

    # 转为padas格式
    O2_pd = pd.DataFrame(O2, index=name)
    CODM_pd = pd.DataFrame(CODM, index=name)
    NH4_pd = pd.DataFrame(NH4, index=name)
    TP_pd = pd.DataFrame(TP, index=name)

    O2_class = O2_pd.copy()
    O2_count = pd.DataFrame()
    for col in range(len(O2_class.columns.values)):
        O2_class[O2_class.columns.values[col]][
            pd.to_numeric(O2_class[O2_class.columns.values[col]], errors='coerce') >= 5] = '优Ⅲ'
        O2_class[O2_class.columns.values[col]][
            pd.to_numeric(O2_class[O2_class.columns.values[col]], errors='coerce') >= 3] = 'Ⅳ'
        O2_class[O2_class.columns.values[col]][
            pd.to_numeric(O2_class[O2_class.columns.values[col]], errors='coerce') < 3] = 'Ⅴ'

        count_3 = ((O2_class[O2_class.columns.values[col]] == '优Ⅲ').sum())
        count_4 = ((O2_class[O2_class.columns.values[col]] == 'Ⅳ').sum())
        count_5 = ((O2_class[O2_class.columns.values[col]] == 'Ⅴ').sum())

        count = pd.Series([count_3, count_4, count_5, 14 - count_3 - count_4 - count_5,
                           format(count_3 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_4 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_5 / (count_3 + count_4 + count_5), '.3f')],
                          index=['优Ⅲ', 'Ⅳ', 'Ⅴ', '异常', '优Ⅲ率', 'Ⅳ率', 'Ⅴ率'])
        O2_count[O2_class.columns.values[col]] = pd.concat([O2_class[O2_class.columns.values[col]], count])

    CODM_class = CODM_pd.copy()
    CODM_count = pd.DataFrame()
    for col in range(len(CODM_class.columns.values)):
        CODM_class[CODM_class.columns.values[col]][
            pd.to_numeric(CODM_class[CODM_class.columns.values[col]], errors='coerce') <= 6] = '优Ⅲ'
        CODM_class[CODM_class.columns.values[col]][
            pd.to_numeric(CODM_class[CODM_class.columns.values[col]], errors='coerce') <= 10] = 'Ⅳ'
        CODM_class[CODM_class.columns.values[col]][
            pd.to_numeric(CODM_class[CODM_class.columns.values[col]], errors='coerce') > 10] = 'Ⅴ'
        count_3 = ((CODM_class[CODM_class.columns.values[col]] == '优Ⅲ').sum())
        count_4 = ((CODM_class[CODM_class.columns.values[col]] == 'Ⅳ').sum())
        count_5 = ((CODM_class[CODM_class.columns.values[col]] == 'Ⅴ').sum())

        count = pd.Series([count_3, count_4, count_5, 14 - count_3 - count_4 - count_5,
                           format(count_3 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_4 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_5 / (count_3 + count_4 + count_5), '.3f')],
                          index=['优Ⅲ', 'Ⅳ', 'Ⅴ', '异常', '优Ⅲ率', 'Ⅳ率', 'Ⅴ率'])
        CODM_count[CODM_class.columns.values[col]] = pd.concat([CODM_class[CODM_class.columns.values[col]], count])

    NH4_class = NH4_pd.copy()
    NH4_count = pd.DataFrame()
    for col in range(len(NH4_class.columns.values)):
        NH4_class[NH4_class.columns.values[col]][
            pd.to_numeric(NH4_class[NH4_class.columns.values[col]], errors='coerce') <= 1.0] = '优Ⅲ'
        NH4_class[NH4_class.columns.values[col]][
            pd.to_numeric(NH4_class[NH4_class.columns.values[col]], errors='coerce') <= 1.5] = 'Ⅳ'
        NH4_class[NH4_class.columns.values[col]][
            pd.to_numeric(NH4_class[NH4_class.columns.values[col]], errors='coerce') > 1.5] = 'Ⅴ'
        count_3 = ((NH4_class[NH4_class.columns.values[col]] == '优Ⅲ').sum())
        count_4 = ((NH4_class[NH4_class.columns.values[col]] == 'Ⅳ').sum())
        count_5 = ((NH4_class[NH4_class.columns.values[col]] == 'Ⅴ').sum())

        count = pd.Series([count_3, count_4, count_5, 14 - count_3 - count_4 - count_5,
                           format(count_3 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_4 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_5 / (count_3 + count_4 + count_5), '.3f')],
                          index=['优Ⅲ', 'Ⅳ', 'Ⅴ', '异常', '优Ⅲ率', 'Ⅳ率', 'Ⅴ率'])
        NH4_count[NH4_class.columns.values[col]] = pd.concat([NH4_class[NH4_class.columns.values[col]], count])

    TP_class = TP_pd.copy()
    kch_TP = (TP_class[TP_class.index == '昆承湖心（省站）']).copy()
    TP_count = pd.DataFrame()
    for col in range(len(TP_class.columns.values)):
        kch_TP[kch_TP.columns.values[col]][
            pd.to_numeric(kch_TP[kch_TP.columns.values[col]], errors='coerce') <= 0.05] = '优Ⅲ'
        kch_TP[kch_TP.columns.values[col]][
            pd.to_numeric(kch_TP[kch_TP.columns.values[col]], errors='coerce') <= 0.1] = 'Ⅳ'
        kch_TP[kch_TP.columns.values[col]][
            pd.to_numeric(kch_TP[kch_TP.columns.values[col]], errors='coerce') > 0.1] = 'Ⅴ'

        TP_class[TP_class.columns.values[col]][
            pd.to_numeric(TP_class[TP_class.columns.values[col]], errors='coerce') <= 0.2] = '优Ⅲ'
        TP_class[TP_class.columns.values[col]][
            pd.to_numeric(TP_class[TP_class.columns.values[col]], errors='coerce') <= 0.3] = 'Ⅳ'
        TP_class[TP_class.columns.values[col]][
            pd.to_numeric(TP_class[TP_class.columns.values[col]], errors='coerce') > 0.3] = 'Ⅴ'

        TP_class.loc['昆承湖心（省站）'][TP_class.columns.values[col]] = kch_TP.loc['昆承湖心（省站）'][kch_TP.columns.values[col]]
        count_3 = ((TP_class[TP_class.columns.values[col]] == '优Ⅲ').sum())
        count_4 = ((TP_class[TP_class.columns.values[col]] == 'Ⅳ').sum())
        count_5 = ((TP_class[TP_class.columns.values[col]] == 'Ⅴ').sum())
        count = pd.Series([count_3, count_4, count_5, 14 - count_3 - count_4 - count_5,
                           format(count_3 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_4 / (count_3 + count_4 + count_5), '.3f'),
                           format(count_5 / (count_3 + count_4 + count_5), '.3f')],
                          index=['优Ⅲ', 'Ⅳ', 'Ⅴ', '异常', '优Ⅲ率', 'Ⅳ率', 'Ⅴ率'])
        TP_count[TP_class.columns.values[col]] = pd.concat([TP_class[TP_class.columns.values[col]], count])

    O2_near7 = O2_count.iloc[0:14, -7:]
    O2_station_near7 = ((O2_near7 == 'Ⅳ') + (O2_near7 == 'Ⅴ')).sum(axis=1)
    CODM_near7 = CODM_count.iloc[0:14, -7:]
    CODM_station_near7 = ((CODM_near7 == 'Ⅳ') + (CODM_near7 == 'Ⅴ')).sum(axis=1)
    NH4_near7 = NH4_count.iloc[0:14, -7:]
    NH4_station_near7 = ((NH4_near7 == 'Ⅳ') + (NH4_near7 == 'Ⅴ')).sum(axis=1)
    TP_near7 = TP_count.iloc[0:14, -7:]
    TP_station_near7 = ((TP_near7 == 'Ⅳ') + (TP_near7 == 'Ⅴ')).sum(axis=1)
    station_near7 = pd.concat([O2_station_near7, CODM_station_near7, NH4_station_near7, TP_station_near7], axis=1)
    station_near7.columns = ['溶解氧', '高锰酸盐指数', '氨氮', '总磷']

    # 写入excel
    writer = pd.ExcelWriter(outxlsx)
    O2_pd.to_excel(writer, sheet_name='溶解氧', index=True, header=True)
    CODM_pd.to_excel(writer, sheet_name='高锰酸盐指数', index=True, header=True)
    NH4_pd.to_excel(writer, sheet_name='氨氮', index=True, header=True)
    TP_pd.to_excel(writer, sheet_name='总磷', index=True, header=True)
    O2_count.to_excel(writer, sheet_name='溶解氧_分级', index=True, header=True)
    CODM_count.to_excel(writer, sheet_name='高锰酸盐指数_分级', index=True, header=True)
    NH4_count.to_excel(writer, sheet_name='氨氮_分级', index=True, header=True)
    TP_count.to_excel(writer, sheet_name='总磷_分级', index=True, header=True)
    station_near7.to_excel(writer, sheet_name='非优Ⅲ类天数', index=True, header=True)
    writer.save()
    station_near7.to_csv(outxlsx1, index=True, header=True)
    return O2_pd.iloc[0:14, -show_days:].apply(pd.to_numeric,errors='coerce'), \
           CODM_pd.iloc[0:14, -show_days:].apply(pd.to_numeric,errors='coerce'), \
           NH4_pd.iloc[0:14, -show_days:].apply(pd.to_numeric,errors='coerce'), \
           TP_pd.iloc[0:14, -show_days:].apply(pd.to_numeric,errors='coerce')


def show_attention(begin_time, idx, parameter,show_days):
    x = pd.date_range(start=begin_time, periods=show_days, freq='d')
    y = idx  # np.random.randint(0, 20, 7)
    # 设置图形显示风格
    plt.style.use('ggplot')
    mpl.rcParams['font.family'] = 'Microsoft YaHei'
    # plt.rcParams['axes.unicode_minus'] = False  # 步骤二（解决坐标轴负数的负号显示问题）

    fig = plt.figure(figsize=(12, 5))
    plt.plot(x, y)
    plt.scatter(x, y, color='black', alpha=0.6)
    plt.fill_between(x, parameter[0], parameter[1], color='#0000FF', label='I类')
    plt.fill_between(x, parameter[1], parameter[2], color='#3b9aff', label='II类')
    plt.fill_between(x, parameter[2], parameter[3], color='#70ffd2', label='III类')
    plt.fill_between(x, parameter[3], parameter[4], color='#e2ff4f', label='IV类')
    plt.fill_between(x, parameter[4], parameter[5], color='#FFA500', label='V类')
    plt.fill_between(x, parameter[5], parameter[6], color='#FF0000', label='劣V类')
    plt.scatter(x, y, color='black', alpha=0.6)
    plt.xlim(x[0], x[-1])
    #plt.xlim(x[0]-datetime.timedelta(hours=8),x[-1])
    plt.ylim(np.nanmin(y) * 0.1, np.nanmax(y) * 1.5)
    for a,b in zip(x,y):
        plt.text(a,b * 1.05,'%.3f'%b, ha='center', va='bottom', fontsize=9)
    # plt.xlabel('日均值')
    plt.ylabel(parameter[7])
    plt.legend(loc=1, bbox_to_anchor=(1.1, 1.01))
    plt.show()


if __name__ == '__main__':
    dirpath = r'E:\Zph\0704常熟日报\1排查资料\日报'
    outxlsx = r'E:\Zph\0704常熟日报\1排查资料\日报汇总.xlsx'
    outxlsx1 = r'E:\Zph\0704常熟日报\1排查资料\非优Ⅲ断面次数.csv'
    attention = {'昆承湖心（省站）': ['CODM', 'TP']
                 #'张桥（省站）': ['O2', 'TP'],
                 #'大义光明村(省站)':['O2','TP','CODM']
                 }
    show_days = 7  #出图显示近n天的数据


    O2_pd_near7, CODM_pd_near7, NH4_pd_near7, TP_pd_near7 = arrange_daily_report(dirpath, outxlsx, outxlsx1,show_days)
    # print(O2_pd_near7)
    begin_time = datetime.datetime.strptime(O2_pd_near7.columns[0], '%Y/%m/%d')
    station_list = list(attention.keys())
    for station in station_list:
        for target in attention[station]:
            if target == 'O2':
                idx = list(map(float, O2_pd_near7.loc[station]))
                parameter = [20, 7.5, 6, 5, 3, 2, 0, 'DO (mg/L)']
                show_attention(begin_time, idx, parameter,show_days)
            elif target == 'CODM':
                idx = list(map(float, CODM_pd_near7.loc[station]))
                parameter = [0, 2, 4, 6, 10, 15, 50, 'CODmn (mg/L)']
                show_attention(begin_time, idx, parameter,show_days)
            elif target == 'NH4':
                idx = list(map(float, NH4_pd_near7.loc[station]))
                parameter = [0, 0.15, 0.5, 1.0, 1.5, 2.0, 10, 'NH4 (mg/L)']
                show_attention(begin_time, idx, parameter,show_days)
            elif target == 'TP':
                idx = list(map(float, TP_pd_near7.loc[station]))
                if station == '昆承湖心（省站）':
                    parameter = [0, 0.01, 0.025, 0.05, 0.1, 0.2, 10, 'TP (mg/L)']
                else:
                    parameter = [0, 0.02, 0.1, 0.2, 0.3, 0.4, 10, 'TP (mg/L)']
                show_attention(begin_time, idx, parameter,show_days)
            else:
                print('error')
